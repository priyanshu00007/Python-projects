import streamlit as st
import pandas as pd
import yagmail
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
import zipfile
import hashlib

# Login system
def login():
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")
        
        if submitted:
            stored_username = "Admin"
            stored_password = hashlib.sha256("Password".encode()).hexdigest()
            
            if username == stored_username and hashlib.sha256(password.encode()).hexdigest() == stored_password:
                return True
            else:
                st.error("Invalid username or password.")
    return False

# Load student data
@st.cache_data
def load_student_data():
    if not os.path.exists('students.xlsx'):
        df = pd.DataFrame({
            "Name": ["John Doe", "Jane Smith", "Mike Johnson", "Emily Brown"],
            "Email": ["john@example.com", "jane@example.com", "mike@example.com", "emily@example.com"],
            "Parent Email": ["parent1@example.com", "parent2@example.com", "parent3@example.com", "parent4@example.com"]
        })
        df.to_excel('students.xlsx', index=False)
    return pd.read_excel('students.xlsx')

# Create absence note
def create_absence_note(student_name, date, time):
    doc = Document()
    doc.add_heading('Student Absence Notification', 0)
    
    doc.add_paragraph(f'Date: {date}')
    doc.add_paragraph(f'Time: {time}')
    doc.add_paragraph(f'Student Name: {student_name}')
    doc.add_paragraph(f'This is to inform you that {student_name} was absent from class on {date} at {time}.')
    doc.add_paragraph('If you have any questions or concerns, please contact the school administration.')
    doc.add_paragraph('Best regards,')
    doc.add_paragraph('School Administration')
    
    return doc

# Send email function
def send_email(to_email, subject, body, attachment=None):
    try:
        sender_email = os.environ.get('SENDER_EMAIL')
        sender_password = os.environ.get('SENDER_PASSWORD')

        if not sender_email or not sender_password:
            st.error("Email configuration is missing. Please set SENDER_EMAIL and SENDER_PASSWORD environment variables.")
            return False

        yag = yagmail.SMTP(sender_email, sender_password)
        contents = [body]
        if attachment:
            contents.append(attachment)
        yag.send(to=to_email, subject=subject, contents=contents)
        return True
    except Exception as e:
        st.error(f"Failed to send email: {str(e)}")
        return False

# Save attendance to Excel
def save_attendance(attendance_df):
    if os.path.exists('attendance.xlsx'):
        existing_df = pd.read_excel('attendance.xlsx')
        updated_df = pd.concat([existing_df, attendance_df], ignore_index=True)
    else:
        updated_df = attendance_df
    updated_df.to_excel('attendance.xlsx', index=False)

# Attendance system
def attendance():
    st.title("Automated Attendance System")
    
    students = load_student_data()

    # Date and time picker for attendance
    col1, col2 = st.columns(2)
    date = col1.date_input("Select Date", datetime.now())
    time = col2.time_input("Select Time", datetime.now())

    # Create checkboxes for attendance
    attendance = {}
    st.write("Mark Attendance:")
    for _, student in students.iterrows():
        attendance[student['Name']] = st.checkbox(student['Name'], value=True)

    if st.button("Submit Attendance"):
        absent_students = [name for name, present in attendance.items() if not present]

        # Record attendance
        attendance_df = pd.DataFrame({
            "Date": [date] * len(students),
            "Time": [time] * len(students),
            "Name": students['Name'],
            "Present": [attendance[name] for name in students['Name']]
        })

        # Save attendance to Excel
        save_attendance(attendance_df)
        st.success("Attendance recorded successfully!")

        # Display attendance
        st.write("Today's Attendance:")
        st.dataframe(attendance_df)

        # Create absence notes and send emails
        if absent_students:
            st.write("Creating absence notes and sending emails...")
            absence_notes = []
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                for name in absent_students:
                    student = students[students['Name'] == name].iloc[0]
                    parent_email = student['Parent Email']
                    
                    # Create absence note
                    doc = create_absence_note(name, date, time)
                    doc_stream = BytesIO()
                    doc.save(doc_stream)
                    doc_stream.seek(0)
                    
                    # Save absence note for download
                    absence_notes.append((f"{name}_absence_note.docx", doc_stream))
                    zip_file.writestr(f"{name}_absence_note.docx", doc_stream.getvalue())
                    
                    # Send email with attachment
                    subject = f"Absence Notification for {name}"
                    body = f"Dear Parent,\n\nPlease find attached the absence notification for {name} for {date} at {time}.\n\nBest regards,\nSchool Administration"
                    if send_email(parent_email, subject, body, (f"{name}_absence_note.docx", doc_stream.getvalue())):
                        st.success(f"Email sent to parent of {name}")
                    else:
                        st.warning(f"Failed to send email for {name}. Please check the email configuration.")

            st.download_button(
                label="Download All Absence Notes",
                data=zip_buffer.getvalue(),
                file_name="absence_notes.zip",
                mime="application/zip"
            )
        else:
            st.success("All students are present today!")

# Streamlit app
def main():
    if 'logged_in' not in st.session_state:
        if login():
            st.session_state.logged_in = True
    elif st.session_state.logged_in:
        attendance()

if __name__ == "__main__":
    main()
# import streamlit as st
# import pandas as pd
# import yagmail
# from datetime import datetime
# import os
# from docx import Document
# from docx.shared import Inches
# from io import BytesIO
# import zipfile
# import hashlib

# # Load student data
# @st.cache_data
# def load_student_data():
#     if not os.path.exists('students.xlsx'):
#         df = pd.DataFrame({
#             "Name": ["John Doe", "Jane Smith", "Mike Johnson", "Emily Brown"],
#             "Email": ["john@example.com", "jane@example.com", "mike@example.com", "emily@example.com"],
#             "Parent Email": ["parent1@example.com", "parent2@example.com", "parent3@example.com", "parent4@example.com"]
#         })
#         df.to_excel('students.xlsx', index=False)
#     return pd.read_excel('students.xlsx')

# # Send email function
# def send_email(to_email, subject, body, attachment=None):
#     try:
#         sender_email = os.environ.get('SENDER_EMAIL')
#         sender_password = os.environ.get('SENDER_PASSWORD')

#         if not sender_email or not sender_password:
#             st.error("Email configuration is missing. Please set SENDER_EMAIL and SENDER_PASSWORD environment variables.")
#             return False

#         yag = yagmail.SMTP(sender_email, sender_password)
#         contents = [body]
#         if attachment:
#             contents.append(attachment)
#         yag.send(to=to_email, subject=subject, contents=contents)
#         return True
#     except Exception as e:
#         st.error(f"Failed to send email: {str(e)}")
#         return False

# # Save attendance to Excel
# def save_attendance(attendance_df):
#     if os.path.exists('attendance.xlsx'):
#         existing_df = pd.read_excel('attendance.xlsx')
#         updated_df = pd.concat([existing_df, attendance_df], ignore_index=True)
#     else:
#         updated_df = attendance_df
#     updated_df.to_excel('attendance.xlsx', index=False)

# # Create absence note
# def create_absence_note(student_name, date, time):
#     doc = Document()
#     doc.add_heading('Student Absence Notification', 0)
    
#     doc.add_paragraph(f'Date: {date}')
#     doc.add_paragraph(f'Time: {time}')
#     doc.add_paragraph(f'Student Name: {student_name}')
#     doc.add_paragraph(f'This is to inform you that {student_name} was absent from class on {date} at {time}.')
#     doc.add_paragraph('If you have any questions or concerns, please contact the school administration.')
#     doc.add_paragraph('Best regards,')
#     doc.add_paragraph('School Administration')
    
#     return doc

# # Hash password
# def hash_password(password):
#     return hashlib.sha256(password.encode()).hexdigest()

# # Check password
# def check_password(stored_password, provided_password):
#     return stored_password == hash_password(provided_password)

# # Login system
# def login():
#     username = st.text_input("Username")
#     password = st.text_input("Password", type="password")
#     if st.button("Login"):
#         stored_username = "Admin"
#         stored_password = hash_password("Password")
        
#         if username == stored_username and check_password(stored_password, password):
#             return True
#         else:
#             st.error("Invalid username or password.")
#     return False

# # Streamlit app
# def main():
#     logged_in = login()
#     if logged_in:
#         st.title("Automated Attendance System")

#         # Load student data
#         students = load_student_data()

#         # Date and time picker for attendance
#         col1, col2 = st.columns(2)
#         date = col1.date_input("Select Date", datetime.now())
#         time = col2.time_input("Select Time", datetime.now())

#         # Create checkboxes for attendance
#         attendance = {}
#         st.write("Mark Attendance:")
#         for _, student in students.iterrows():
#             attendance[student['Name']] = st.checkbox(student['Name'], value=True)

#         if st.button("Submit Attendance"):
#             absent_students = [name for name, present in attendance.items() if not present]

#             # Record attendance
#             attendance_df = pd.DataFrame({
#                 "Date": [date] * len(students),
#                 "Time": [time] * len(students),
#                 "Name": students['Name'],
#                 "Present": [attendance[name] for name in students['Name']]
#             })

#             # Save attendance to Excel
#             save_attendance(attendance_df)
#             st.success("Attendance recorded successfully!")

#                         # Display attendance
#             st.write("Today's Attendance:")
#             st.dataframe(attendance_df)

#             # Create absence notes and send emails
#             if absent_students:
#                 st.write("Creating absence notes and sending emails...")
#                 absence_notes = []
#                 zip_buffer = BytesIO()
#                 with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
#                     for name in absent_students:
#                         student = students[students['Name'] == name].iloc[0]
#                         parent_email = student['Parent Email']
                        
#                         # Create absence note
#                         doc = create_absence_note(name, date, time)
#                         doc_stream = BytesIO()
#                         doc.save(doc_stream)
#                         doc_stream.seek(0)
                        
#                         # Save absence note for download
#                         absence_notes.append((f"{name}_absence_note.docx", doc_stream))
#                         zip_file.writestr(f"{name}_absence_note.docx", doc_stream.getvalue())
                        
#                         # Send email with attachment
#                         subject = f"Absence Notification for {name}"
#                         body = f"Dear Parent,\n\nPlease find attached the absence notification for {name} for {date} at {time}.\n\nBest regards,\nSchool Administration"
#                         if send_email(parent_email, subject, body, (f"{name}_absence_note.docx", doc_stream.getvalue())):
#                             st.success(f"Email sent to parent of {name}")
#                         else:
#                             st.warning(f"Failed to send email for {name}. Please check the email configuration.")

#                 st.download_button(
#                     label="Download All Absence Notes",
#                     data=zip_buffer.getvalue(),
#                     file_name="absence_notes.zip",
#                     mime="application/zip"
#                 )
#             else:
#                 st.success("All students are present today!")

# if __name__ == "__main__":
#     main()
# import streamlit as st
# import pandas as pd
# import yagmail
# from datetime import datetime
# import os
# from docx import Document
# from docx.shared import Inches
# from io import BytesIO
# import zipfile

# # Load student data
# @st.cache_data
# def load_student_data():
#     if not os.path.exists('students.xlsx'):
#         df = pd.DataFrame({
#             "Name": ["John Doe", "Jane Smith", "Mike Johnson", "Emily Brown"],
#             "Email": ["john@example.com", "jane@example.com", "mike@example.com", "emily@example.com"],
#             "Parent Email": ["parent1@example.com", "parent2@example.com", "parent3@example.com", "parent4@example.com"]
#         })
#         df.to_excel('students.xlsx', index=False)
#     return pd.read_excel('students.xlsx')

# # Send email function
# def send_email(to_email, subject, body, attachment=None):
#     try:
#         sender_email = os.environ.get('SENDER_EMAIL')
#         sender_password = os.environ.get('SENDER_PASSWORD')

#         if not sender_email or not sender_password:
#             st.error("Email configuration is missing. Please set SENDER_EMAIL and SENDER_PASSWORD environment variables.")
#             return False

#         yag = yagmail.SMTP(sender_email, sender_password)
#         contents = [body]
#         if attachment:
#             contents.append(attachment)
#         yag.send(to=to_email, subject=subject, contents=contents)
#         return True
#     except Exception as e:
#         st.error(f"Failed to send email: {str(e)}")
#         return False

# # Save attendance to Excel
# def save_attendance(attendance_df):
#     if os.path.exists('attendance.xlsx'):
#         existing_df = pd.read_excel('attendance.xlsx')
#         updated_df = pd.concat([existing_df, attendance_df], ignore_index=True)
#     else:
#         updated_df = attendance_df
#     updated_df.to_excel('attendance.xlsx', index=False)

# # Create absence note
# def create_absence_note(student_name, date, time):
#     doc = Document()
#     doc.add_heading('Student Absence Notification', 0)
    
#     doc.add_paragraph(f'Date: {date}')
#     doc.add_paragraph(f'Time: {time}')
#     doc.add_paragraph(f'Student Name: {student_name}')
#     doc.add_paragraph(f'This is to inform you that {student_name} was absent from class on {date} at {time}.')
#     doc.add_paragraph('If you have any questions or concerns, please contact the school administration.')
#     doc.add_paragraph('Best regards,')
#     doc.add_paragraph('School Administration')
    
#     return doc

# # Streamlit app
# def main():
#     st.title("Automated Attendance System")

#     # Load student data
#     students = load_student_data()

#     # Date and time picker for attendance
#     col1, col2 = st.columns(2)
#     date = col1.date_input("Select Date", datetime.now())
#     time = col2.time_input("Select Time", datetime.now())

#     # Create checkboxes for attendance
#     attendance = {}
#     st.write("Mark Attendance:")
#     for _, student in students.iterrows():
#         attendance[student['Name']] = st.checkbox(student['Name'], value=True)

#     if st.button("Submit Attendance"):
#         absent_students = [name for name, present in attendance.items() if not present]

#         # Record attendance
#         attendance_df = pd.DataFrame({
#             "Date": [date] * len(students),
#             "Time": [time] * len(students),
#             "Name": students['Name'],
#             "Present": [attendance[name] for name in students['Name']]
#         })

#         # Save attendance to Excel
#         save_attendance(attendance_df)
#         st.success("Attendance recorded successfully!")

#         # Display attendance
#         st.write("Today's Attendance:")
#         st.dataframe(attendance_df)

#         # Create absence notes and send emails
#         if absent_students:
#             st.write("Creating absence notes and sending emails...")
#             absence_notes = []
#             for name in absent_students:
#                 student = students[students['Name'] == name].iloc[0]
#                 parent_email = student['Parent Email']
                
#                 # Create absence note
#                 doc = create_absence_note(name, date, time)
#                 doc_stream = BytesIO()
#                 doc.save(doc_stream)
#                 doc_stream.seek(0)
                
#                 # Save absence note for download
#                 absence_notes.append((f"{name}_absence_note.docx", doc_stream))
                
#                 # Send email with attachment
#                 subject = f"Absence Notification for {name}"
#                 body = f"Dear Parent,\n\nPlease find attached the absence notification for {name} for {date} at {time}.\n\nBest regards,\nSchool Administration"
#                 if send_email(parent_email, subject, body, (f"{name}_absence_note.docx", doc_stream.getvalue())):
#                                         st.success(f"Email sent to parent of {name}")
#                 else:
#                     st.warning(f"Failed to send email for {name}. Please check the email configuration.")

#             # Create a zip file of all absence notes
#             zip_buffer = BytesIO()
#             with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
#                 for file_name, data in absence_notes:
#                     zip_file.writestr(file_name, data.getvalue())
            
#             st.download_button(
#                 label="Download All Absence Notes",
#                 data=zip_buffer.getvalue(),
#                 file_name="absence_notes.zip",
#                 mime="application/zip"
#             )
#         else:
#             st.success("All students are present today!")

# if __name__ == "__main__":
#     main()